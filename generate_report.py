# -*- coding: utf-8 -*-
"""
Script tạo báo cáo ứng dụng Trợ Lý AI Học Tập (BAO CAO APP - Updated).
Dựa trên cấu trúc file gốc, cập nhật nội dung cho đúng app hiện tại.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ===== STYLE SETUP =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Title style
title_style = doc.styles['Title']
title_style.font.name = 'Times New Roman'
title_style.font.size = Pt(16)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(15)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

# ===== HELPER FUNCTIONS =====
def add_paragraph(text, style_name='Normal', bold=False, italic=False, indent=False):
    p = doc.add_paragraph(text, style=style_name)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Paragraph')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.7)
    return p


# ====================================================================
# TITLE
# ====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BÁO CÁO ỨNG DỤNG\nTRỢ LÝ AI HỌC TẬP')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Ứng dụng trí tuệ nhân tạo hỗ trợ học sinh giải bài tập bằng hình ảnh')
run2.italic = True
run2.font.size = Pt(13)
run2.font.name = 'Times New Roman'
run2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

# Author info
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author.paragraph_format.space_before = Pt(12)
run_a = author.add_run('Tác giả: Trần Thị Kim Thoa\nTrường THPT Hoàng Diệu\nĐịa chỉ: Số 1 Mạc Đĩnh Chi, phường Phú Lợi, thành phố Cần Thơ')
run_a.font.size = Pt(12)
run_a.font.name = 'Times New Roman'

doc.add_paragraph()  # spacer

# ====================================================================
# PHẦN I. ĐẶT VẤN ĐỀ
# ====================================================================
doc.add_heading('PHẦN I. ĐẶT VẤN ĐỀ', level=1)

doc.add_heading('1. Lý do chọn đề tài', level=2)
add_paragraph(
    'Trong bối cảnh chuyển đổi số giáo dục, học sinh ngày nay cần ' 
    'những công cụ học tập thông minh hỗ trợ giải bài tập mọi lúc, mọi nơi. '
    'Việc học tập với gia sư truyền thống tốn nhiều thời gian, chi phí cao, '
    'và phụ thuộc vào lịch hẹn. Trong khi đó, các mô hình AI hiện đại như '
    'Google Gemini có khả năng phân tích hình ảnh và trả lời câu hỏi '
    'với độ chính xác ngày càng cao.'
)
add_paragraph(
    'Vì vậy, ứng dụng "Trợ Lý AI Học Tập" ra đời nhằm giúp học sinh '
    'chỉ cần chụp ảnh bài tập, AI sẽ tự động phân tích đề bài và đưa ra '
    'lời giải chi tiết từng bước, bao gồm cả công thức toán học hiển thị '
    'trực quan. Ứng dụng hoạt động 24/7, miễn phí, và có thể sử dụng '
    'trên mọi thiết bị có trình duyệt web — đặc biệt tối ưu cho điện thoại thông minh.'
)

doc.add_heading('2. Mục đích nghiên cứu', level=2)
add_bullet('Xây dựng ứng dụng Web App tích hợp AI nhằm:')
add_bullet('Giúp học sinh giải bài tập mọi lúc mọi nơi chỉ bằng cách chụp ảnh đề bài.', level=1)
add_bullet('Hiển thị lời giải chi tiết từng bước, dễ hiểu, kèm công thức toán học trực quan (MathJax/LaTeX).', level=1)
add_bullet('Hỗ trợ tính năng đọc lời giảng bằng giọng nói (Text-to-Speech) để học sinh có thể "nghe giảng" mọi lúc.', level=1)
add_bullet('Cung cấp giao diện thân thiện, tối ưu cho smartphone, giúp học sinh dễ dàng sử dụng.', level=1)

doc.add_heading('3. Đối tượng và phạm vi nghiên cứu', level=2)
add_bullet('Đối tượng: Học sinh THPT và THCS cần hỗ trợ giải bài tập các môn Toán, Lý, Hóa.')
add_bullet('Công nghệ: Google Gemini API (gemini-3-flash-preview, gemini-3-pro-preview, gemini-2.5-flash), ReactJS, TailwindCSS, Vite, Framer Motion, MathJax v3.')
add_bullet('Phạm vi: Ứng dụng web đáp ứng (responsive), tối ưu cho smartphone, có thể triển khai trên Vercel.')

# ====================================================================
# PHẦN II. NỘI DUNG
# ====================================================================
doc.add_heading('PHẦN II. NỘI DUNG', level=1)

doc.add_heading('1. Cơ sở lý luận', level=2)
add_paragraph(
    'Ứng dụng được xây dựng dựa trên kiến trúc Modern Web App (SPA – Single Page Application) '
    'kết hợp với Generative AI. Sử dụng kỹ thuật Prompt Engineering để tinh chỉnh câu trả lời '
    'của AI cho phù hợp với bối cảnh giáo dục Việt Nam. Hệ thống tận dụng API Vision của Google Gemini '
    'để phân tích hình ảnh bài tập và tự động trích xuất lời giải có cấu trúc (JSON), '
    'sau đó hiển thị trực quan trên giao diện.'
)
add_paragraph(
    'Thiết kế giao diện tuân theo nguyên tắc "Mobile-first" với phong cách Claymorphism '
    '(bo góc mềm 16-24px, viền dày 3px, hiệu ứng đổ bóng 3D mềm mại) — được tham chiếu từ '
    'bộ kỹ năng UI/UX Pro Max chuyên nghiệp. Font chữ Baloo 2 được chọn vì tính thân thiện, '
    'dễ đọc, phù hợp ứng dụng giáo dục.'
)

doc.add_heading('2. Các tính năng nổi bật (Giải pháp)', level=2)
add_paragraph(
    'Hệ thống "Trợ Lý AI Học Tập" cung cấp các giải pháp toàn diện cho việc học tập:'
)

doc.add_heading('a. Chụp/Tải ảnh bài tập:', level=3)
add_bullet('Học sinh có thể chụp ảnh trực tiếp từ camera hoặc tải ảnh từ thư viện ảnh trên điện thoại.')
add_bullet('Hỗ trợ các định dạng JPG, PNG với dung lượng tối đa 10MB.')
add_bullet('Giao diện xem trước ảnh để xác nhận trước khi gửi cho AI phân tích.')

doc.add_heading('b. AI Giải bài tập tự động:', level=3)
add_bullet('Sử dụng Google Gemini (mô hình ngôn ngữ lớn) để phân tích hình ảnh bài tập.')
add_bullet('Trả về lời giải có cấu trúc gồm: tiêu đề bài toán, các bước giải chi tiết, và kết quả cuối cùng.')
add_bullet('Hỗ trợ cơ chế Fallback tự động — nếu một mô hình AI gặp lỗi hoặc hết quota, hệ thống tự động chuyển sang mô hình dự phòng (gemini-3-flash → gemini-3-pro → gemini-2.5-flash).')

doc.add_heading('c. Hiển thị công thức toán học trực quan (MathJax):', level=3)
add_bullet('Tích hợp MathJax v3 để hiển thị công thức Toán/Lý/Hóa chuẩn LaTeX ngay trong lời giải.')
add_bullet('Hỗ trợ cả inline math ($...$) và display math ($$...$$).')
add_bullet('Công thức hiển thị rõ ràng trên mọi kích thước màn hình, đặc biệt tối ưu cho điện thoại.')

doc.add_heading('d. Đọc lời giảng bằng giọng nói (Text-to-Speech):', level=3)
add_bullet('Tính năng "Nghe giảng" cho phép AI đọc lại lời giải bằng giọng nói tiếng Việt.')
add_bullet('Giúp học sinh vừa nhìn bài giải vừa nghe giải thích — phù hợp nhiều phong cách học.')

doc.add_heading('e. Quản lý API Key và Mô hình AI:', level=3)
add_bullet('Giao diện cài đặt dạng Bottom Sheet (trượt từ dưới lên, phù hợp thao tác một tay).')
add_bullet('Cho phép người dùng chọn mô hình AI ưu tiên: Gemini 3 Flash (mặc định), Gemini 3 Pro, hoặc Gemini 2.5 Flash.')
add_bullet('API Key được lưu trữ an toàn trong localStorage của trình duyệt.')
add_bullet('Hướng dẫn lấy API Key miễn phí từ Google AI Studio ngay trong ứng dụng.')

doc.add_heading('f. Giao diện hiện đại tối ưu cho smartphone:', level=3)
add_bullet('Thiết kế theo phong cách Claymorphism — mềm mại, bo góc tròn, hiệu ứng 3D nhẹ nhàng.')
add_bullet('Touch target tối thiểu 48px theo chuẩn Mobile Design.')
add_bullet('Font chữ Baloo 2 thân thiện, dễ đọc.')
add_bullet('Animation mượt mà với Framer Motion (hiệu ứng chuyển trang, nút bấm, loading).')
add_bullet('Hỗ trợ iPhone safe area (notch, home indicator).')
add_bullet('Hiển thị thông tin tác giả với avatar và logo trường học.')

# ====================================================================
# PHẦN III. KIẾN TRÚC HỆ THỐNG
# ====================================================================
doc.add_heading('PHẦN III. KIẾN TRÚC HỆ THỐNG', level=1)

doc.add_heading('1. Công nghệ sử dụng', level=2)

# Technology table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
# Header row
hdr = table.rows[0]
for i, text in enumerate(['Thành phần', 'Công nghệ', 'Mục đích']):
    cell = hdr.cells[i]
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(12)

rows_data = [
    ['Frontend', 'ReactJS + TypeScript', 'Xây dựng giao diện SPA'],
    ['Styling', 'TailwindCSS', 'Hệ thống CSS utility-first'],
    ['Design System', 'Claymorphism + Baloo 2', 'Phong cách UI bo tròn, thân thiện'],
    ['Animation', 'Framer Motion', 'Hiệu ứng chuyển trang, nút bấm'],
    ['AI Engine', 'Google Gemini API', 'Phân tích hình ảnh, giải bài tập'],
    ['Math Rendering', 'MathJax v3', 'Hiển thị công thức LaTeX trực quan'],
    ['Text-to-Speech', 'Web Speech API', 'Đọc lời giải bằng giọng nói'],
    ['Build Tool', 'Vite', 'Bundler nhanh cho phát triển và production'],
    ['Deployment', 'Vercel', 'Hosting và tự động deploy'],
]

for row_data in rows_data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(12)

doc.add_paragraph()

doc.add_heading('2. Cấu trúc thư mục dự án', level=2)
structure = """trolyhoctap/
├── index.html              (Trang HTML chính + MathJax CDN)
├── vercel.json             (Cấu hình SPA routing cho Vercel)
├── GV/                     (Thông tin tác giả)
│   ├── avatar.jpg
│   ├── logo.jpg
│   └── THÔNG TIN GV.txt
└── src/
    ├── index.css           (Hệ thống thiết kế Claymorphism)
    ├── main.tsx            (Entry point)
    ├── App.tsx             (Component chính, quản lý flow)
    ├── components/
    │   ├── Header.tsx          (Thanh tiêu đề + nút Cài đặt)
    │   ├── AuthorProfile.tsx   (Thẻ thông tin tác giả)
    │   ├── UploadSection.tsx   (Khu vực chụp/tải ảnh)
    │   ├── LoadingView.tsx     (Màn hình chờ AI phân tích)
    │   ├── SolutionCard.tsx    (Hiển thị lời giải + MathJax)
    │   ├── MathContent.tsx     (Component render LaTeX)
    │   ├── ComparisonTable.tsx (Bảng so sánh Gia sư vs AI)
    │   └── SettingsModal.tsx   (Modal cài đặt API Key/Model)
    └── services/
        └── gemini.ts           (Kết nối Gemini API + Fallback)"""

p_struct = doc.add_paragraph()
run_s = p_struct.add_run(structure)
run_s.font.name = 'Consolas'
run_s.font.size = Pt(10)

# ====================================================================
# PHẦN IV. KẾT QUẢ VÀ HƯỚNG PHÁT TRIỂN
# ====================================================================
doc.add_heading('PHẦN IV. KẾT QUẢ VÀ HƯỚNG PHÁT TRIỂN', level=1)

doc.add_heading('1. Kết quả đạt được', level=2)
add_bullet('Ứng dụng hoạt động ổn định trên mọi trình duyệt hiện đại (Chrome, Safari, Firefox), đặc biệt trên smartphone.')
add_bullet('AI phân tích hình ảnh bài tập và trả lời chính xác với lời giải từng bước rõ ràng.')
add_bullet('Công thức toán học hiển thị đẹp, trực quan nhờ MathJax — giải quyết triệt để vấn đề hiển thị LaTeX trên web.')
add_bullet('Cơ chế Fallback tự động giữa 3 mô hình AI đảm bảo ứng dụng luôn hoạt động ngay cả khi một model hết quota.')
add_bullet('Giao diện Claymorphism thân thiện, chuyên nghiệp, tối ưu cho thao tác một tay trên điện thoại.')
add_bullet('Thời gian phản hồi trung bình dưới 10 giây cho mỗi bài tập.')
add_bullet('Tính năng Text-to-Speech hỗ trợ học sinh nghe lại lời giải bằng tiếng Việt.')

doc.add_heading('2. Hướng phát triển', level=2)
add_bullet('Phát triển phiên bản Mobile App native (React Native) để tối ưu hơn trên iOS/Android.')
add_bullet('Tích hợp tính năng nhận diện giọng nói (Voice-to-Text) để ra lệnh cho trợ lý.')
add_bullet('Thêm hệ thống lưu trữ lịch sử bài tập đã giải để học sinh ôn tập.')
add_bullet('Hỗ trợ thêm nhiều môn học: Văn, Anh, Sinh, Sử, Địa.')
add_bullet('Xây dựng tính năng "Bài tập tương tự" — AI tự động đề xuất bài luyện thêm sau khi giải xong.')
add_bullet('Tích hợp bảng xếp hạng và thành tích để tạo động lực học tập.')

# ====================================================================
# SAVE
# ====================================================================
output_path = 'BAO CAO APP - Updated.docx'
doc.save(output_path)
print(f'✅ Đã tạo file báo cáo mới: {output_path}')
print(f'📄 Kích thước: {os.path.getsize(output_path)} bytes')
